import os
import json
import streamlit as st
from openai import OpenAI
from docx import Document
from docx.shared import Pt

# =========================================================
# Helpers
# =========================================================
def get_client():
    api_key = st.secrets.get("OPENAI_API_KEY", None) or os.getenv("OPENAI_API_KEY")
    if not api_key:
        st.error("Missing OPENAI_API_KEY. Add it in Streamlit Secrets.")
        st.stop()
    return OpenAI(api_key=api_key)

def get_model():
    return st.secrets.get("OPENAI_MODEL", "gpt-4o-mini")

def clean_text(x: str) -> str:
    return (x or "").strip()

def build_inputs(payload: dict) -> str:
    return json.dumps(payload, ensure_ascii=False, indent=2)

DISCOVERY_SYSTEM = """You are a senior management consultant.
Your job: convert messy discovery notes into a premium, neutral, executive-ready "Discovery Intelligence Report".
Rules:
- Do NOT provide solutions or recommendations.
- Do NOT propose vendors, tools, or implementation steps.
- Be factual and neutral; avoid blame.
- If information is missing, label it as "Unknown" or "Not confirmed".
- Use crisp consulting language, but keep it understandable.
- Prefer structured outputs (headings, bullets, matrices, short tables).
- If user content contains sensitive details, do not invent names or specifics.
"""

DISCOVERY_USER_TEMPLATE = """Create a Discovery Intelligence Report from the inputs.

OUTPUT FORMAT (Markdown):
1. Title block (client name if provided; otherwise "Client"; date placeholder; meeting type)
2. Executive Narrative Map
   - Problem–Pressure–Consequence narrative (6–10 lines)
   - Why now (3 bullets)
   - What happens if not done (3 bullets)
3. Scope & Objective Clarity
   - Project objective (as heard)
   - In-scope / Out-of-scope (based on notes only; if unknown say Unknown)
4. Stakeholder & Power Reality Map
   - Stakeholder list by role/department
   - Influence vs Ownership matrix (ASCII table)
   - Accountability vs Authority mismatch signals (bullets)
5. KPI & Load Signal Snapshot
   - KPI burden signals
   - BAU impact signals
   - Evidence statements (quote-like paraphrases, neutral; 3–6 items)
6. Organizational Context Timeline (last 3–5 years if possible)
   - CEO transitions, org changes, M&A, vendor history, change programs
   - If missing, write "Not provided"
7. Risk Exposure Canvas (no solutions)
   - Strategic / Operational / Financial / Organizational / Cultural
   - For each: risk statement + trigger + consequence
8. Engagement Justification (Why external support is logical)
   - Constraint vs Neutrality table (ASCII)
9. Open Questions & Data Needed
   - Grouped list (Governance, KPIs, Stakeholders, Contracts, Change history, Budget)
10. Meeting Close Summary (3 bullets)
   - What we heard
   - What we did today (discovery)
   - Proposed next step (ONLY: "alignment workshop / diagnostic deep-dive" style, not a solution)

INPUTS:
{inputs_json}
"""

def generate_report(payload: dict) -> str:
    client = get_client()
    model = get_model()
    inputs_json = build_inputs(payload)
    user_prompt = DISCOVERY_USER_TEMPLATE.format(inputs_json=inputs_json)

    resp = client.responses.create(
        model=model,
        input=[
            {"role": "system", "content": DISCOVERY_SYSTEM},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0.3,
    )
    return resp.output_text

def markdown_to_docx(md_text: str) -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for raw in md_text.splitlines():
        line = raw.rstrip()

        if not line.strip():
            doc.add_paragraph("")
            continue

        if line.startswith("### "):
            doc.add_heading(line.replace("### ", ""), level=3)
        elif line.startswith("## "):
            doc.add_heading(line.replace("## ", ""), level=2)
        elif line.startswith("# "):
            doc.add_heading(line.replace("# ", ""), level=1)
        elif line.lstrip().startswith(("- ", "* ")):
            doc.add_paragraph(line.lstrip()[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)

    import io
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

# =========================================================
# Session init
# =========================================================
st.set_page_config(page_title="Discovery Intelligence Report Generator", layout="wide")

st.session_state.setdefault("report_md", "")
st.session_state.setdefault("docx_bytes", None)
st.session_state.setdefault("last_error", "")

# =========================================================
# Payload builder (reads current session_state)
# =========================================================
def build_payload_from_state():
    include_open_questions = st.session_state.get("include_open_questions", True)
    payload = {
        "client_name": clean_text(st.session_state.get("client_name", "")) or "Client",
        "meeting_type": st.session_state.get("meeting_type", "Discovery / Intake"),
        "project_name": clean_text(st.session_state.get("project_name", "")),
        "transcript_or_notes": clean_text(st.session_state.get("transcript", "")),
        "structured_inputs": {
            "project_objective": clean_text(st.session_state.get("objective", "")),
            "why_initiated_problem_trigger": clean_text(st.session_state.get("why_now", "")),
            "benefiting_departments": clean_text(st.session_state.get("beneficiaries", "")),
            "impacted_people": clean_text(st.session_state.get("impacted_people", "")),
            "kpi_burden": clean_text(st.session_state.get("kpis", "")),
            "if_not_done_consequences": clean_text(st.session_state.get("constraints_if_not_done", "")),
            "internal_challenges": clean_text(st.session_state.get("internal_challenges", "")),
            "org_changes": clean_text(st.session_state.get("org_changes", "")),
            "ceo_info": clean_text(st.session_state.get("ceo_info", "")),
            "previous_ceo_problems": clean_text(st.session_state.get("prior_ceo_issues", "")),
            "why_external_vendor": clean_text(st.session_state.get("vendor_reason", "")),
            "why_not_listening_internally": clean_text(st.session_state.get("listening_issue", "")),
            "ownership_and_misalignment": clean_text(st.session_state.get("ownership_misalignment", "")),
            "contracts_dependencies": clean_text(st.session_state.get("contracts", "")),
            "ma_and_culture": clean_text(st.session_state.get("ma_history", "")),
            "budget_duration_payment": clean_text(st.session_state.get("budget_duration_payment", "")),
            "long_term_vision_and_next": clean_text(st.session_state.get("long_term", "")),
        },
        "report_constraints": {
            "no_solutions": True,
            "include_open_questions": include_open_questions,
        }
    }
    return payload

# =========================================================
# Generation (called from Output tab)
# =========================================================
def run_generation_from_output():
    try:
        payload = build_payload_from_state()

        if not payload["transcript_or_notes"] and all(not v for v in payload["structured_inputs"].values()):
            st.session_state["last_error"] = "Please paste at least a transcript/notes OR fill at least one structured field."
            st.session_state["report_md"] = ""
            st.session_state["docx_bytes"] = None
            return

        st.session_state["last_error"] = ""

        report_md = generate_report(payload)
        st.session_state["report_md"] = report_md

        if st.session_state.get("include_docx", True):
            st.session_state["docx_bytes"] = markdown_to_docx(report_md)
        else:
            st.session_state["docx_bytes"] = None

    except Exception as e:
        st.session_state["last_error"] = f"Generation failed: {e}"
        st.session_state["report_md"] = ""
        st.session_state["docx_bytes"] = None

def clear_report():
    st.session_state["report_md"] = ""
    st.session_state["docx_bytes"] = None
    st.session_state["last_error"] = ""

# =========================================================
# UI
# =========================================================
st.title("Discovery Intelligence Report Generator (Streamlit)")
st.caption("Turn messy discovery notes into a premium executive-ready report (no solutions).")

with st.sidebar:
    st.header("Settings")
    st.write(f"Model: `{get_model()}`")
    st.slider("Tone (Neutral ↔ Strong)", 0, 10, 3, key="tone")
    st.divider()
    st.subheader("Quality checks")
    st.checkbox("Include 'Open Questions & Data Needed' section", value=True, key="include_open_questions")
    st.checkbox("Enable DOCX download", value=True, key="include_docx")

tab1, tab2 = st.tabs(["Input", "Output"])

# -----------------------------
# Input tab (only collects inputs)
# -----------------------------
with tab1:
    st.subheader("Client & meeting info")
    colA, colB, colC = st.columns(3)
    with colA:
        st.text_input("Client name (optional)", value="", key="client_name")
    with colB:
        st.selectbox(
            "Meeting type",
            ["Discovery / Intake", "Stakeholder interview", "Project kick-off (discovery)", "Other"],
            index=0,
            key="meeting_type",
        )
    with colC:
        st.text_input("Project name (optional)", value="", key="project_name")

    st.subheader("Paste transcript / rough notes")
    st.text_area(
        "Transcript / notes (paste here)",
        height=220,
        placeholder="Paste meeting transcript or your rough notes here...",
        key="transcript",
    )

    st.subheader("Structured inputs (optional but improves output)")
    col1, col2 = st.columns(2)

    with col1:
        st.text_area("Project objective (as stated)", height=90, key="objective")
        st.text_area("Why initiated / what problem triggered it", height=90, key="why_now")
        st.text_area("Departments that benefit", height=90, key="beneficiaries")
        st.text_area("People impacted (counts/roles)", height=90, key="impacted_people")

    with col2:
        st.text_area("KPI burden (how many / examples / pain)", height=90, key="kpis")
        st.text_area("What happens if project not done", height=90, key="constraints_if_not_done")
        st.text_area("Challenges solving internally", height=90, key="internal_challenges")
        st.text_area("Recent organizational changes / leadership", height=90, key="org_changes")

    st.subheader("Governance / history / commercials (optional)")
    col3, col4 = st.columns(2)

    with col3:
        st.text_area("CEO info (join date, style, priorities)", height=90, key="ceo_info")
        st.text_area("Problems faced by previous CEO", height=90, key="prior_ceo_issues")
        st.text_area("Why external vendor is required", height=90, key="vendor_reason")
        st.text_area("Why internal management not listening", height=90, key="listening_issue")

    with col4:
        st.text_area("Ownership / misalignment issues", height=90, key="ownership_misalignment")
        st.text_area("Contracts ending / dependencies", height=90, key="contracts")
        st.text_area("M&A / cultural integration issues", height=90, key="ma_history")
        st.text_area("Budget, duration, payment method", height=90, key="budget_duration_payment")

    st.text_area("Long-term company vision / what happens after project", height=90, key="long_term")

    st.info("Go to the Output tab and click **Generate/Refresh** to create the report.")

# -----------------------------
# Output tab (Generate/Refresh lives here)
# -----------------------------
with tab2:
    st.subheader("Generated report")

    colX, colY, colZ = st.columns([1.2, 1.0, 2.8])
    with colX:
        st.button("Generate / Refresh report", type="primary", on_click=run_generation_from_output)
    with colY:
        st.button("Clear report", on_click=clear_report)
    with colZ:
        st.caption("Tip: update inputs in the Input tab, then click **Generate/Refresh** here.")

    if st.session_state.get("last_error"):
        st.error(st.session_state["last_error"])

    report_md = st.session_state.get("report_md", "")
    if not report_md:
        st.info("No report yet. Click **Generate / Refresh report** above.")
    else:
        st.markdown(report_md)

        colD, colE = st.columns(2)
        with colD:
            st.download_button(
                "Download Markdown",
                data=report_md.encode("utf-8"),
                file_name="discovery_intelligence_report.md",
                mime="text/markdown",
            )
        with colE:
            if st.session_state.get("docx_bytes"):
                st.download_button(
                    "Download DOCX",
                    data=st.session_state["docx_bytes"],
                    file_name="discovery_intelligence_report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
